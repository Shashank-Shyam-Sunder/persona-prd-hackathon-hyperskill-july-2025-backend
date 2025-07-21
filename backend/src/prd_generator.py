import os
import time
import pandas as pd
from typing import List, Tuple
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches # For potentially adding more advanced styling

from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import PromptTemplate
from .persona_config import PERSONA_TO_FOLDER, PERSONA_DISPLAY_NAMES

# ───────────────────────────────────────────────
# 🔐 Load environment variables
# ───────────────────────────────────────────────
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    raise ValueError("GOOGLE_API_KEY not found in environment variables.")

# ───────────────────────────────────────────────
# 🧠 Gemini model initialization
# ───────────────────────────────────────────────
# Note: Ensure you have your Google API key configured for Gemini access
model = ChatGoogleGenerativeAI(
    model="gemini-1.5-flash", # Or "gemini-1.5-pro" if you prefer higher quality/longer context (and have access)
    google_api_key=GOOGLE_API_KEY
)

# ───────────────────────────────────────────────
# 📄 Prompt template (Final – Updated with hallucination safeguards)
# ───────────────────────────────────────────────
PRD_PROMPT_TEMPLATE = """
You are a senior Product Manager AI assistant at PersonaPRD, an AI-powered product discovery platform.

Your task is to generate a structured Product Requirements Document (PRD) draft based on the following pain point summaries collected from user community data.

**Persona:** {persona_name}

**Pain Points:**
{pain_points}

**Instructions:**

Write a PRD draft that includes the following 6 items in this exact order:

1. A **single bold title line** like:
   **PRD Draft: [short, product-focused name for the solution]**
   (e.g., "**PRD Draft: AI-Powered BI Onboarding and Analytics Platform**")

2. **Problem Summary:** A clear summary of the problem(s) these pain points represent. Use simple, direct language.

3. **Why This Problem Matters:** Explain why this problem is significant specifically for the **{persona_name}** persona. Highlight the impact on productivity, workflows, or business goals.

4. **Potential Solution Overview:** Provide a concise solution concept that addresses these pain points.

5. **Suggested MVP Features:** List 3–5 minimum viable product features as bullet points, phrased as actionable features (e.g. “Feature X does Y to solve Z”).

6. **Next Steps:** Outline immediate steps the team should take to validate and build this solution (e.g. user interviews, prototype, sprint planning).

**Strict output rules:**
- Do NOT include any metadata like dates, authors, product IDs, or document codes.
- Do NOT add extra sections beyond the 6 listed above.
- Do NOT mention that this is AI-generated.
- Do NOT include headings like "Product Requirements Document" or "PersonaPRD PRD".

**Tone guidelines:**
- Use clear, professional, and confident language.
- Avoid generic filler phrases or vague platitudes.
- Tailor the writing to a product team preparing for sprint planning.

PRD Draft:
"""

prompt = PromptTemplate(
    input_variables=["pain_points", "persona_name"],
    template=PRD_PROMPT_TEMPLATE,
)

prd_chain = prompt | model

# ───────────────────────────────────────────────
# 🔧 Core Functions (Modified for API Integration)
# ───────────────────────────────────────────────

def generate_prd_content( # Renamed from generate_prd
    pain_point_summaries: List[str],
    persona_display_name: str,
    subreddit_file: str,
    selected_cluster_ids: List[int], # These should now be 0-indexed as per model
    num_posts: int
) -> str:
    """
    Generates the text content of the PRD using the LLM.
    Includes a header with source information.
    """
    # Adjust cluster_ids for display if they are 0-indexed internally (add 1 for human readability)
    display_cluster_ids = [str(cid + 1) for cid in selected_cluster_ids]

    header = (
        f"**📌 Source: `{subreddit_file}`**\n"
        f"**📊 Clusters analyzed:** {', '.join(display_cluster_ids)} "
        f"(based on {num_posts} posts)\n\n"
        f"**📂 Pain Points Selected:**\n"
        + "\n".join([f"- {pp}" for pp in pain_point_summaries]) + "\n\n"
        f"💡 *To review raw discussions behind these pain points, open `cluster_visualization.html` in your browser.*\n"
        f"This file is located in the same folder as this PRD.\n\n"
        f"---\n"
    )

    combined_pain_points = "\n".join([f"- {pp}" for pp in pain_point_summaries])
    prd_body = prd_chain.invoke({"pain_points": combined_pain_points, "persona_name": persona_display_name}).content.strip()
    return header + prd_body


def load_summaries(persona_key: str, subreddit_file: str) -> pd.DataFrame:
    """
    Loads the pain point summaries from the processed data directory.
    """
    persona_folder = PERSONA_TO_FOLDER.get(persona_key)
    if not persona_folder:
        raise ValueError(f"Unknown persona key: {persona_key}. Please ensure it's in persona_config.PERSONA_TO_FOLDER.")

    subreddit_folder = subreddit_file.replace(".json", "")
    summaries_path = os.path.abspath(os.path.join(
        os.path.dirname(__file__), "..", "data", "processed", persona_folder, subreddit_folder, "pain_point_summaries.csv"
    ))
    
    if not os.path.exists(summaries_path):
        raise FileNotFoundError(
            f"Pain point summaries not found at expected path: {summaries_path}. "
            f"Please ensure the main clustering pipeline has been run successfully for persona '{persona_key}' and subreddit '{subreddit_file}'."
        )
    df = pd.read_csv(summaries_path)
    # Ensure cluster_id column exists and is integer type
    if 'cluster_id' not in df.columns:
        raise ValueError("The 'pain_point_summaries.csv' file is missing the 'cluster_id' column.")
    df['cluster_id'] = df['cluster_id'].astype(int)
    return df


def save_prd_to_docx(prd_text: str, output_path: str) -> str: # Added return type
    """
    Saves the generated PRD text content to a DOCX file.
    Includes basic formatting for bold and list items.
    """
    doc = Document()
    # Add a main title for the document itself (distinct from the PRD's internal title)
    doc.add_heading("Generated Product Requirements Document", level=0) 
    doc.add_paragraph("This document was automatically generated by the PersonaPRD AI Assistant.")
    doc.add_paragraph("") # Blank line

    for line in prd_text.split("\n"):
        line_stripped = line.strip()
        if line_stripped:
            if line_stripped.startswith('**PRD Draft:'):
                doc.add_heading(line_stripped.replace('**', ''), level=1)
            elif line_stripped.startswith('**') and line_stripped.endswith('**'):
                # For bold section titles like 'Problem Summary:'
                doc.add_heading(line_stripped.replace('**', ''), level=2)
            elif line_stripped.startswith('- '):
                # For bullet points
                doc.add_paragraph(line_stripped.lstrip('- ').strip(), style='List Bullet')
            elif line_stripped == '---': # Separator line
                doc.add_paragraph('---')
            else:
                doc.add_paragraph(line_stripped)
    
    doc.save(output_path)
    print(f"\n✅ PRD saved to: {output_path}")
    return output_path # Crucial: Return the path for the API to send back

# ───────────────────────────────────────────────
# 💡 New Orchestrator Function for API Integration
# ───────────────────────────────────────────────
def generate_prd_for_api(
    persona_key: str,
    subreddit_file: str,
    selected_cluster_ids: List[int] # These are expected to be 0-indexed from frontend
) -> str: # Returns the absolute path to the generated PRD DOCX file
    """
    Orchestrates the PRD generation process for use with a web API.
    Loads summaries, filters based on user selection, generates PRD content using LLM,
    and saves the PRD to a DOCX file in the processed data directory.
    """
    print(f"[{time.ctime()}] Loading summaries for persona: {persona_key}, subreddit: {subreddit_file}")
    # 1. Load all summaries for the given persona and subreddit
    all_summaries_df = load_summaries(persona_key, subreddit_file)

    # 2. Filter summaries based on selected cluster IDs
    # Ensure selected_cluster_ids are treated as integers for comparison
    selected_cluster_ids_int = [int(cid) for cid in selected_cluster_ids]
    filtered_df = all_summaries_df[all_summaries_df["cluster_id"].isin(selected_cluster_ids_int)]

    if filtered_df.empty:
        # Also return a ValueError to signal that no data was found for selection
        raise ValueError(
            f"No valid pain point summaries found for selected cluster IDs {selected_cluster_ids_int} "
            f"for persona '{persona_key}' and subreddit '{subreddit_file}'. "
            f"Available cluster IDs: {sorted(all_summaries_df['cluster_id'].tolist()) if not all_summaries_df.empty else 'None'}"
        )

    pain_point_summaries_list = filtered_df["pain_point_summary"].tolist()
    
    # Get the display name for the persona
    persona_display_name = PERSONA_DISPLAY_NAMES.get(persona_key, persona_key)

    num_posts_in_selected_clusters = filtered_df["num_posts"].sum() if "num_posts" in filtered_df.columns else 0

    print(f"[{time.ctime()}] Generating PRD content for {len(pain_point_summaries_list)} selected pain points...")
    # 3. Generate PRD content using the LLM
    prd_content = generate_prd_content(
        pain_point_summaries_list,
        persona_display_name,
        subreddit_file,
        selected_cluster_ids, # Pass original 0-indexed list to content generator
        num_posts_in_selected_clusters
    )

    # 4. Determine output path for the PRD (same directory as other processed outputs)
    persona_folder = PERSONA_TO_FOLDER.get(persona_key)
    subreddit_folder = subreddit_file.replace(".json", "")
    
    # This must match the PROCESSED_DATA_DIR logic in api/main.py
    output_dir_base = os.path.abspath(os.path.join(
        os.path.dirname(__file__), "..", "data", "processed", persona_folder, subreddit_folder
    ))
    os.makedirs(output_dir_base, exist_ok=True) # Ensure the specific output directory exists

    timestamp = pd.Timestamp.now().strftime("%Y%m%d_%H%M%S")
    
    # Attempt to extract actual title from generated content for a better filename
    prd_filename_base = "Generated_PRD" # Default in case title extraction fails
    if prd_content.strip().startswith('**PRD Draft:'):
        # Extract actual title from generated content for a better filename
        generated_title_line = prd_content.split('\n')[0]
        # Remove bold markers and "PRD Draft: " prefix
        actual_title = generated_title_line.replace('**PRD Draft: ', '').replace('**', '').strip()
        # Sanitize for filename
        prd_filename_base = "".join(c for c in actual_title if c.isalnum() or c in (' ', '-', '_')).strip()
        prd_filename_base = prd_filename_base.replace(' ', '_')[:50] # Limit length and replace spaces

    prd_filename = f"{prd_filename_base}_PRD_{persona_key}_{subreddit_folder}_{timestamp}.docx"
    prd_output_path = os.path.join(output_dir_base, prd_filename)

    print(f"[{time.ctime()}] Saving PRD to: {prd_output_path}")
    # 5. Save the PRD to a DOCX file
    saved_path = save_prd_to_docx(prd_content, prd_output_path)
    return saved_path # Return the absolute path to the saved DOCX file