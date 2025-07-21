# src/utils.py

from docx import Document

def save_prd_as_docx(prd_text: str, filepath: str):
    """
    Save the generated PRD text into a nicely formatted .docx file.

    Args:
        prd_text (str): Text of the PRD draft.
        filepath (str): Where to save the document.
    """
    doc = Document()
    doc.add_heading("Product Requirements Document (PRD)", level=1)

    for line in prd_text.strip().split("\n"):
        stripped = line.strip()

        if stripped.startswith("1.") or stripped.startswith("2.") or stripped.startswith("3.") or \
           stripped.startswith("4.") or stripped.startswith("5."):
            doc.add_heading(stripped, level=2)
        elif stripped.startswith("-"):
            doc.add_paragraph(stripped, style='List Bullet')
        else:
            doc.add_paragraph(stripped)

    doc.save(filepath)
