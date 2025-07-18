# src/prd_generator.py

import os
import pandas as pd
from typing import List
from dotenv import load_dotenv
from docx import Document

from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import PromptTemplate
from .persona_config import PERSONA_TO_FOLDER, PERSONA_DISPLAY_NAMES

# ───────────────────────────────────────────────
# 🔐 Load environment variables
# ───────────────────────────────────────────────
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    raise ValueError("GOOGLE_API_KEY not found in environment variables.")

# ───────────────────────────────────────────────
# 🧠 Gemini model initialization
# ───────────────────────────────────────────────
model = ChatGoogleGenerativeAI(
    model="gemini-1.5-flash",
    google_api_key=GOOGLE_API_KEY
)

# ───────────────────────────────────────────────
# 📄 Prompt template (Final – Updated with hallucination safeguards)
# ───────────────────────────────────────────────
PRD_PROMPT_TEMPLATE = """
You are a senior Product Manager AI assistant at PersonaPRD, an AI-powered product discovery platform.

Your task is to generate a structured Product Requirements Document (PRD) draft based on the following pain point summaries collected from user community data.

**Persona:** {persona_name}

**Pain Points:**
{pain_points}

**Instructions:**

Write a PRD draft that includes the following 6 items in this exact order:

1. A **single bold title line** like:
   **PRD Draft: [short, product-focused name for the solution]**
   (e.g., "**PRD Draft: AI-Powered BI Onboarding and Analytics Platform**")

2. **Problem Summary:** A clear summary of the problem(s) these pain points represent. Use simple, direct language.

3. **Why This Problem Matters:** Explain why this problem is significant specifically for the **{persona_name}** persona. Highlight the impact on productivity, workflows, or business goals.

4. **Potential Solution Overview:** Provide a concise solution concept that addresses these pain points.

5. **Suggested MVP Features:** List 3–5 minimum viable product features as bullet points, phrased as actionable features (e.g. “Feature X does Y to solve Z”).

6. **Next Steps:** Outline immediate steps the team should take to validate and build this solution (e.g. user interviews, prototype, sprint planning).

**Strict output rules:**
- Do NOT include any metadata like dates, authors, product IDs, or document codes.
- Do NOT add extra sections beyond the 6 listed above.
- Do NOT mention that this is AI-generated.
- Do NOT include headings like "Product Requirements Document" or "PersonaPRD PRD".

**Tone guidelines:**
- Use clear, professional, and confident language.
- Avoid generic filler phrases or vague platitudes.
- Tailor the writing to a product team preparing for sprint planning.

PRD Draft:
"""


prompt = PromptTemplate(
    input_variables=["pain_points", "persona_name"],
    template=PRD_PROMPT_TEMPLATE,
)

prd_chain = prompt | model

# ───────────────────────────────────────────────
# 🔧 Core Functions
# ───────────────────────────────────────────────

def generate_prd(
    pain_point_summaries: List[str],
    persona_display_name: str,
    subreddit_file: str,
    selected_cluster_ids: List[int],
    num_posts: int
) -> str:
    header = (
        f"**📌 Source: `{subreddit_file}`**\n"
        f"**📊 Clusters analyzed:** {', '.join(str(cid+1) for cid in selected_cluster_ids)} "
        f"(based on {num_posts} posts)\n\n"
        f"**📂 Pain Points Selected:**\n"
        + "\n".join([f"- {pp}" for pp in pain_point_summaries]) + "\n\n"
        f"💡 *To review raw discussions behind these pain points, open `cluster_visualization.html` in your browser.*\n"
        f"This file is located in the same folder as this PRD.\n\n"
        f"---\n"
    )

    combined = "\n".join([f"- {pp}" for pp in pain_point_summaries])
    prd_body = prd_chain.invoke({"pain_points": combined, "persona_name": persona_display_name}).content.strip()
    return header + prd_body


def load_summaries(persona_key: str, subreddit_file: str) -> pd.DataFrame:
    persona_folder = PERSONA_TO_FOLDER.get(persona_key)
    if not persona_folder:
        raise ValueError(f"Unknown persona key: {persona_key}")

    subreddit_folder = subreddit_file.replace(".json", "")
    summaries_path = os.path.abspath(os.path.join(
        os.path.dirname(__file__), "..", "data", "processed", persona_folder, subreddit_folder, "pain_point_summaries.csv"
    ))
    if not os.path.exists(summaries_path):
        raise FileNotFoundError(f"Pain point summaries not found at {summaries_path}")
    return pd.read_csv(summaries_path)


def select_pain_points(df: pd.DataFrame) -> (List[str], List[int], int):
    print("\n📌 Found", len(df), "pain point summaries. Here they are:\n")
    for idx, row in df.iterrows():
        print(f"[{row['cluster_id']+1}] {row['pain_point_summary']}\n")

    selected = input("🔍 Enter comma-separated cluster IDs to include in PRD (e.g. 1,3,6): ")
    cluster_ids = [int(cid.strip())-1 for cid in selected.split(",") if cid.strip().isdigit() and 1 <= int(cid.strip()) <= 10]
    selected_df = df[df["cluster_id"].isin(cluster_ids)]
    summaries = selected_df["pain_point_summary"].tolist()
    cluster_ids_used = list(selected_df["cluster_id"].unique())
    num_posts = selected_df["num_posts"].sum() if "num_posts" in selected_df.columns else 0
    return summaries, cluster_ids_used, num_posts


def save_prd_to_docx(prd_text: str, output_path: str):
    doc = Document()
    doc.add_heading("Product Requirements Document (Generated)", level=1)
    for line in prd_text.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())
    doc.save(output_path)
    print(f"\n✅ PRD saved to: {output_path}")
